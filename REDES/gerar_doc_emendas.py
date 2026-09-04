from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Paleta ────────────────────────────────────────────────────────────────────
AZUL_ESCURO  = RGBColor(0x1A, 0x37, 0x6C)   # azul PP
AZUL_MEDIO   = RGBColor(0x1E, 0x5F, 0xAD)
AZUL_CLARO   = RGBColor(0xD6, 0xE4, 0xF7)
LARANJA      = RGBColor(0xE8, 0x6A, 0x10)
CINZA_TEXTO  = RGBColor(0x3A, 0x3A, 0x3A)
CINZA_SUAVE  = RGBColor(0xF2, 0xF4, 0xF8)
BRANCO       = RGBColor(0xFF, 0xFF, 0xFF)
VERDE        = RGBColor(0x1E, 0x8B, 0x4C)

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = str(color)  # RGBColor.__str__ retorna hex sem #
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val", "single"))
            el.set(qn("w:sz"),    str(val.get("sz", 4)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color=AZUL_ESCURO, size=20, bold=True, space_before=16, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold       = bold
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    return p

def add_normal(doc, text, size=10.5, color=CINZA_TEXTO, italic=False, space_after=4, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.bold       = bold
    run.italic     = italic
    run.font.name  = "Calibri"
    return p

def add_label_value(doc, label, value, label_color=AZUL_MEDIO, value_color=CINZA_TEXTO):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + " ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = label_color
    r1.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = value_color
    r2.font.name = "Calibri"
    return p

def add_angle_box(doc, text):
    """Caixa azul claro com o ângulo de post."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, AZUL_CLARO)
    set_cell_border(cell,
        left={"val": "single", "sz": 12, "color": "1E5FAD"},
    )
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("Ângulo de post: ")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = AZUL_ESCURO
    r.font.name = "Calibri"
    r2 = p.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = AZUL_ESCURO
    r2.italic = True
    r2.font.name = "Calibri"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 72)
    run.font.color.rgb = RGBColor(0xCC, 0xD6, 0xE8)
    run.font.size = Pt(7)
    run.font.name = "Calibri"

def add_category_header(doc, emoji, title, subtitle=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, AZUL_ESCURO)
    cell.width = Cm(16.5)
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(f"{emoji}  {title}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BRANCO
    r.font.name = "Calibri"
    if subtitle:
        p2 = cell.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.4)
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xAA, 0xC4, 0xE8)
        r2.font.name = "Calibri"
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_emenda_card(doc, numero, titulo, valor, status, municipio, descricao, angulo=None, destaque=False):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    bg = CINZA_SUAVE if not destaque else RGBColor(0xFF, 0xF3, 0xE6)
    set_cell_bg(cell, bg)
    set_cell_border(cell, left={"val":"single","sz":18,"color": "E86A10" if destaque else "1E5FAD"})
    cell.width = Cm(16.5)

    # Cabeçalho da card
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.left_indent  = Cm(0.4)
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after  = Pt(1)

    # Número e título
    r_num = p_title.add_run(f"{numero}. ")
    r_num.bold = True
    r_num.font.size = Pt(11)
    r_num.font.color.rgb = LARANJA if destaque else AZUL_MEDIO
    r_num.font.name = "Calibri"
    r_title = p_title.add_run(titulo)
    r_title.bold = True
    r_title.font.size = Pt(11)
    r_title.font.color.rgb = AZUL_ESCURO
    r_title.font.name = "Calibri"

    if destaque:
        r_star = p_title.add_run("  ★ DESTAQUE")
        r_star.bold = True
        r_star.font.size = Pt(8)
        r_star.font.color.rgb = LARANJA
        r_star.font.name = "Calibri"

    # Valor / Status / Município em linha
    p_meta = cell.add_paragraph()
    p_meta.paragraph_format.left_indent = Cm(0.4)
    p_meta.paragraph_format.space_after = Pt(3)
    for label, val in [("Valor:", valor), ("Status:", status), ("Município:", municipio)]:
        r_l = p_meta.add_run(label + " ")
        r_l.bold = True
        r_l.font.size = Pt(9)
        r_l.font.color.rgb = AZUL_MEDIO
        r_l.font.name = "Calibri"
        r_v = p_meta.add_run(val + "    ")
        r_v.font.size = Pt(9)
        r_v.font.color.rgb = CINZA_TEXTO
        r_v.font.name = "Calibri"

    # Descrição
    p_desc = cell.add_paragraph()
    p_desc.paragraph_format.left_indent = Cm(0.4)
    p_desc.paragraph_format.space_after = Pt(3 if not angulo else 5)
    r_desc = p_desc.add_run(descricao)
    r_desc.font.size = Pt(9.5)
    r_desc.font.color.rgb = CINZA_TEXTO
    r_desc.font.name = "Calibri"

    # Ângulo dentro da card
    if angulo:
        p_ang = cell.add_paragraph()
        p_ang.paragraph_format.left_indent = Cm(0.4)
        p_ang.paragraph_format.space_after = Pt(6)
        r_al = p_ang.add_run("Ângulo de post: ")
        r_al.bold = True
        r_al.font.size = Pt(9)
        r_al.font.color.rgb = LARANJA if destaque else AZUL_MEDIO
        r_al.font.name = "Calibri"
        r_at = p_ang.add_run(angulo)
        r_at.font.size = Pt(9)
        r_at.italic = True
        r_at.font.color.rgb = AZUL_ESCURO
        r_at.font.name = "Calibri"
    else:
        cell.paragraphs[-1].paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── CAPA ──────────────────────────────────────────────────────────────────────
# Faixa de título
tbl_capa = doc.add_table(rows=1, cols=1)
tbl_capa.alignment = WD_TABLE_ALIGNMENT.LEFT
c = tbl_capa.cell(0,0)
set_cell_bg(c, AZUL_ESCURO)
p_capa = c.paragraphs[0]
p_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_capa.paragraph_format.space_before = Pt(18)
p_capa.paragraph_format.space_after  = Pt(6)
r = p_capa.add_run("DEP. PEPÊ COLLAÇO")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BRANCO; r.font.name = "Calibri"

p2 = c.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run("Emendas com Maior Potencial para Redes Sociais")
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0xAA,0xC4,0xE8); r2.font.name = "Calibri"

p3 = c.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(14)
r3 = p3.add_run("Seleção estratégica por valor, impacto visual e potencial de engajamento")
r3.font.size = Pt(10); r3.italic = True; r3.font.color.rgb = RGBColor(0x88,0xAA,0xCC); r3.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Subtítulo intro
add_normal(doc,
    "Este documento reúne as emendas do mandato com maior potencial de gerar conteúdo "
    "autêntico e engajador nas redes sociais. Cada emenda foi analisada pelo ângulo de "
    "valor, impacto humano, apelo visual e resultado concreto.",
    size=10.5, color=CINZA_TEXTO, space_after=8)

add_separator(doc)

# ─── RANKING TOP 5 ─────────────────────────────────────────────────────────────
add_heading(doc, "🏆  TOP 5 — Comece por Aqui", size=14, color=LARANJA, space_before=12, space_after=6)

ranking = [
    ("1º", "Parque Infantil Adaptado para TEA — Tubarão",
     "Emocional, único, alcance orgânico explosivo. Raridade no Brasil."),
    ("2º", "Ginásio de Esportes — Sangão (R$ 5,9 milhões)",
     "Maior obra individual. Já PRONTO. Foto/vídeo de inauguração com a comunidade."),
    ("3º", "Creche Joelma de Bona — Paulo Lopes (R$ 2,3 milhões)",
     "Nome real + criança + resultado concreto = narrativa perfeita."),
    ("4º", "Ambulância SAMU — Gravatal (R$ 480 mil)",
     "Simples, poderoso, literalmente salva vidas. Entrega muito fotogênica."),
    ("5º", "Dragagem Porto de Laguna / Molhes (R$ 5 milhões)",
     "Histórico, grandioso. Identidade ribeirinha da região como pano de fundo."),
]

tbl_rank = doc.add_table(rows=len(ranking), cols=2)
tbl_rank.alignment = WD_TABLE_ALIGNMENT.LEFT
col_widths = [Cm(2), Cm(14)]

for i, (pos, titulo, desc) in enumerate(ranking):
    row = tbl_rank.rows[i]
    row.cells[0].width = col_widths[0]
    row.cells[1].width = col_widths[1]
    set_cell_bg(row.cells[0], LARANJA)
    set_cell_bg(row.cells[1], CINZA_SUAVE if i % 2 == 0 else BRANCO)

    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(8)
    r0 = p0.add_run(pos)
    r0.bold = True; r0.font.size = Pt(13); r0.font.color.rgb = BRANCO; r0.font.name = "Calibri"

    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.left_indent  = Cm(0.3)
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(1)
    rt = p1.add_run(titulo)
    rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = AZUL_ESCURO; rt.font.name = "Calibri"

    p2r = row.cells[1].add_paragraph()
    p2r.paragraph_format.left_indent = Cm(0.3)
    p2r.paragraph_format.space_after = Pt(6)
    rd = p2r.add_run(desc)
    rd.font.size = Pt(9); rd.font.color.rgb = CINZA_TEXTO; rd.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(8)
add_separator(doc)

# ─── CATEGORIA 1 — OBRAS ESTRUTURANTES ─────────────────────────────────────────
add_category_header(doc, "🏗️", "CATEGORIA 1 — OBRAS ESTRUTURANTES",
    "Grande impacto visual | Fotos de obra | Inauguração com comunidade")

add_emenda_card(doc, 1, "Ginásio de Esportes — Sangão",
    "R$ 5.900.000", "PRONTO", "Sangão",
    "A maior obra individual do mandato. Um ginásio completo para uma cidade pequena — esporte, eventos, shows, formaturas. "
    "Impacto na qualidade de vida de toda a população. Conteúdo de inauguração com comunidade reunida é ouro nas redes.",
    angulo='"O maior ginásio que Sangão já teve. E chegou com a assinatura do Pepê."',
    destaque=True)

add_emenda_card(doc, 2, "Creche Joelma de Bona — Paulo Lopes",
    "R$ 2.300.000", "PRONTO", "Paulo Lopes",
    "Creche que leva um nome de pessoa real — abre espaço para contar a história por trás do nome, entrevista com famílias, "
    "crianças atendidas. Muito emotivo e visualmente rico. Já entregue.",
    angulo='"Uma creche que tem nome, tem história, tem cara. As crianças de Paulo Lopes agora têm um lugar digno para começar."',
    destaque=True)

add_emenda_card(doc, 3, "Dragagem do Porto de Laguna / Molhes",
    "R$ 5.000.000", "PUBLICADO", "Laguna/Tubarão (Bancada)",
    "Projeto executivo para obra histórica da região portuária. Laguna tem identidade ribeirinha fortíssima. "
    "A dragagem é pauta antiga que nunca saía do papel. Conta a história de uma região que vive do mar.",
    angulo='"Laguna é porto. O projeto que vai mudar a história do Sul Catarinense finalmente saiu do papel."')

add_emenda_card(doc, 4, "Abertura da Av. Pedro Zapelini — Tubarão",
    "R$ 2.000.000", "PUBLICADO", "Tubarão (Bancada)",
    "Obra de mobilidade urbana no município-sede do mandato. Ótima para mostrar que Tubarão está na prioridade. "
    "Antes e depois de avenida nova é conteúdo de alto impacto.",
    angulo='"Tubarão vai ganhar uma nova avenida. E o Pepê garantiu os recursos para isso acontecer."')

add_emenda_card(doc, 5, "Pavimentação Av. Eliete de Souza — Pescaria Brava",
    "R$ 2.000.000", "PAGO", "Pescaria Brava",
    "Obra de grande porte já paga e concluída. Conteúdo de resultado concreto: 'prometemos, entregamos.' "
    "Ideal para série de posts com antes e depois das ruas.",
    angulo='Série "ENTREGUE" — foto da rua asfaltada com moradores. Resultado concreto, sem enrolação.')

add_emenda_card(doc, 6, "Infraestrutura Viária — Pescaria Brava",
    "R$ 2.316.000", "PAGO", "Pescaria Brava (Bancada 2025)",
    "Emenda de bancada de mais de R$ 2 milhões, já paga. Município pequeno com obra de altíssimo impacto. "
    "Potencial de engajamento local enorme. Mobiliza toda a comunidade.")

add_separator(doc)

# ─── CATEGORIA 2 — TEA / AUTISMO ───────────────────────────────────────────────
add_category_header(doc, "💙", "CATEGORIA 2 — TEA / AUTISMO",
    "Bandeira central do mandato | Alto apelo emocional | Engajamento orgânico enorme")

add_emenda_card(doc, 7, "Parque Infantil Adaptado para Pessoas com TEA — Tubarão",
    "R$ 100.000", "PUBLICADO", "Tubarão",
    "A emenda com MAIOR POTENCIAL EMOCIONAL de todo o mandato. Um parque infantil pensado especificamente para "
    "crianças com autismo — raríssimo no Brasil. Narrativa de inclusão, família e pertencimento. "
    "Vídeo de criança com TEA brincando pela primeira vez em espaço adaptado é conteúdo que viraliza.",
    angulo='"Toda criança merece brincar. As crianças com autismo de Tubarão agora têm um espaço pensado para elas."',
    destaque=True)

add_emenda_card(doc, 8, "Equoterapia e Terapia Ocupacional TEA — Múltiplos Municípios",
    "R$ 100 mil a R$ 400 mil por município", "Vários status", "Laguna, Gravatal, Braço do Norte, São Ludgero, Armazém, Pedras Grandes, Orleans, Jacinto Machado e outros",
    "Série de posts 'Em cada cidade, uma família recebe tratamento.' Mostrar o mapa de SC com os municípios atendidos. "
    "Humaniza o mandato e demonstra abrangência geográfica de uma bandeira real.",
    angulo='"Autismo não tem fronteira. O mandato do Pepê chega onde o tratamento precisa estar."')

add_emenda_card(doc, 9, "Salas Sensoriais na Rede de Ensino — Treviso e Nova Veneza",
    "R$ 100.000 cada", "AGUARDANDO PORTARIA", "Treviso e Nova Veneza",
    "Sala sensorial dentro de escola pública = inclusão educacional concreta. Tema pouco explorado politicamente "
    "mas com enorme engajamento de pais, educadores e profissionais de saúde.",
    angulo='"Escola que inclui de verdade. Salas sensoriais chegando em SC graças ao mandato Pepê Collaço."')

add_emenda_card(doc, 10, "Custeio do Programa de Atendimento TEA — Jaguaruna",
    "R$ 150.000", "AGUARDANDO DOCUMENTAÇÃO", "Jaguaruna",
    "Comunidade custeando serviço especializado para autistas. Depoimento de profissional ou família local = "
    "conteúdo autêntico e humanizado.")

add_separator(doc)

# ─── CATEGORIA 3 — SAÚDE ───────────────────────────────────────────────────────
add_category_header(doc, "🏥", "CATEGORIA 3 — SAÚDE",
    "Hospitais filantrópicos | Ambulâncias | Exames que antes eram longe")

add_emenda_card(doc, 11, "Ambulância SAMU — Gravatal",
    "R$ 480.000", "PRONTO", "Gravatal",
    "Ambulância nova literalmente salva vidas. A imagem de entrega com equipe do SAMU, uniforme, faixa e "
    "viatura nova é muito impactante. Simples, direto, emocional.",
    angulo='"Uma ambulância nova em Gravatal. Porque saúde de qualidade é direito de quem mora no interior."',
    destaque=True)

add_emenda_card(doc, 12, "Equipamentos de Endoscopia — Jacinto Machado",
    "R$ 400.000", "PUBLICADO", "Jacinto Machado",
    "Município pequeno ganhando equipamento de diagnóstico que antes obrigava o paciente a viajar 2h. "
    "História de quem não precisa mais ir longe para fazer exame.",
    angulo='"Agora o morador de Jacinto Machado não precisa ir longe para fazer endoscopia. O recurso veio do Pepê."')

add_emenda_card(doc, 13, "Tomografias e Ressonâncias para o SUS — Garopaba",
    "R$ 300.000", "PAGO", "Garopaba",
    "Fila do SUS para exames de imagem em cidade com população crescente. Resultado concreto e já pago.")

add_emenda_card(doc, 14, "Consultas de Oftalmologia no SUS — Criciúma",
    "R$ 400.000", "AGUARDANDO DOCUMENTAÇÃO", "Criciúma",
    "Visão é tema universal. Idosos que voltam a enxergar. Emoção fácil de traduzir em vídeo curto.",
    angulo='"Muita gente perdeu a visão esperando na fila. Com o Pepê, essa fila está sendo zerada."')

add_emenda_card(doc, 15, "Hospital de Caridade de Jaguaruna",
    "R$ 400.000", "EM ANÁLISE", "Jaguaruna",
    "Hospital filantrópico pequeno atendendo região inteira. Narrativa de 'hospital que não pode fechar' "
    "ressoa muito com comunidades do interior.")

add_separator(doc)

# ─── CATEGORIA 4 — EDUCAÇÃO ─────────────────────────────────────────────────────
add_category_header(doc, "🏫", "CATEGORIA 4 — EDUCAÇÃO E INFÂNCIA",
    "Creches | Ginásios | Projetos para jovens")

add_emenda_card(doc, 16, "Construção de CEI (Creche) — Bairro Ouro Negro, Forquilhinha",
    "R$ 900.000", "AGUARDANDO PORTARIA", "Forquilhinha",
    "Nova creche para o Bairro Ouro Negro. O nome do bairro já conta uma história. "
    "Crianças pequenas = empatia automática. Alto potencial de mobilização local.",
    angulo='"O Ouro Negro de Forquilhinha vai ter a creche que merece."')

add_emenda_card(doc, 17, "Cobertura de Cancha — Centro de Eventos, Paulo Lopes",
    "R$ 900.000", "PUBLICADO", "Paulo Lopes",
    "Centro de eventos coberto = shows, festas, formaturas, casamentos. Cidade pequena ganhando espaço de eventos "
    "mobiliza toda a comunidade. Potencial de viralizar localmente.")

add_emenda_card(doc, 18, "Revitalização do Ginásio + Pavimentação de Ruas — Paulo Lopes",
    "R$ 600.000", "PRONTO", "Paulo Lopes",
    "Resultado concreto duplo — esporte e mobilidade na mesma cidade. PRONTO = conteúdo de resultado para postar já.")

add_separator(doc)

# ─── CATEGORIA 5 — SEGURANÇA ───────────────────────────────────────────────────
add_category_header(doc, "🚔", "CATEGORIA 5 — SEGURANÇA PÚBLICA",
    "Bombeiros | Polícia | Câmeras de segurança")

add_emenda_card(doc, 19, "Viatura 4x4 para o Corpo de Bombeiros — Tubarão",
    "R$ 360.000", "PAGO", "Tubarão",
    "Bombeiros são universalmente amados. Imagem de entrega da viatura com a equipe fardada = foto/vídeo poderoso. "
    "Mensagem de segurança que une toda a população.",
    angulo='"Quando o Corpo de Bombeiros tem o equipamento certo, mais vidas são salvas."')

add_emenda_card(doc, 20, "Sistema de Câmeras de Segurança — Biguaçu",
    "R$ 250.000", "AGUARDANDO PORTARIA", "Biguaçu",
    "Tema que gera engajamento em qualquer cidade. Medo da violência urbana é preocupação universal. "
    "Demonstra que o mandato tem abrangência estadual.")

add_emenda_card(doc, 21, "Viatura para a Delegacia — Laguna",
    "R$ 150.000", "PRONTO", "Laguna",
    "Delegada recebendo viatura. Imagem forte. Reforça compromisso com segurança pública na região.")

add_separator(doc)

# ─── CATEGORIA 6 — RURAL E TECNOLOGIA ──────────────────────────────────────────
add_category_header(doc, "🌾", "CATEGORIA 6 — RURAL E TECNOLOGIA",
    "Agricultor | Interior | Inovação no campo")

add_emenda_card(doc, 22, "Drone Agrícola — Treze de Maio",
    "R$ 100.000", "AGUARDANDO PORTARIA", "Treze de Maio",
    "Tecnologia no campo: agricultor usando drone para monitorar lavoura. Imagem inusitada e moderna. "
    "Gera curiosidade e quebra a narrativa de que interior é atrasado.",
    angulo='"O agricultor de Treze de Maio vai usar drone. Santa Catarina é assim."',
    destaque=True)

add_emenda_card(doc, 23, "2 Tratores Agrícolas — Braço do Norte",
    "R$ 400.000", "PRONTO", "Braço do Norte",
    "Tratores entregues para o pequeno agricultor. Imagem clássica e emocionante do Sul Catarinense. "
    "Já PRONTO — conteúdo de resultado imediato.")

add_separator(doc)

# ─── TABELA RESUMO ──────────────────────────────────────────────────────────────
add_heading(doc, "📋  Tabela Resumo — Todas as Emendas Selecionadas", size=13, color=AZUL_ESCURO, space_before=14)

headers = ["#", "Emenda", "Município", "Valor", "Status"]
rows_data = [
    ("1",  "Ginásio de Esportes",               "Sangão",           "R$ 5.900.000", "PRONTO"),
    ("2",  "Creche Joelma de Bona",              "Paulo Lopes",      "R$ 2.300.000", "PRONTO"),
    ("3",  "Dragagem Porto de Laguna",           "Laguna/Tubarão",   "R$ 5.000.000", "PUBLICADO"),
    ("4",  "Abertura Av. Pedro Zapelini",        "Tubarão",          "R$ 2.000.000", "PUBLICADO"),
    ("5",  "Pavimentação Av. Eliete de Souza",  "Pescaria Brava",   "R$ 2.000.000", "PAGO"),
    ("6",  "Infraestrutura Viária (Bancada)",    "Pescaria Brava",   "R$ 2.316.000", "PAGO"),
    ("7",  "Parque Infantil Adaptado TEA",       "Tubarão",          "R$ 100.000",   "PUBLICADO"),
    ("8",  "Equoterapia TEA (múltiplos)",        "Vários municípios","Vários",       "Vários"),
    ("9",  "Salas Sensoriais na Rede Pública",   "Treviso / N. Veneza","R$ 100k cada","AG. PORTARIA"),
    ("10", "Programa TEA Jaguaruna",             "Jaguaruna",        "R$ 150.000",   "AG. DOC."),
    ("11", "Ambulância SAMU",                    "Gravatal",         "R$ 480.000",   "PRONTO"),
    ("12", "Equipamentos Endoscopia",            "Jacinto Machado",  "R$ 400.000",   "PUBLICADO"),
    ("13", "Tomografias e Ressonâncias",         "Garopaba",         "R$ 300.000",   "PAGO"),
    ("14", "Oftalmologia SUS",                   "Criciúma",         "R$ 400.000",   "AG. DOC."),
    ("15", "Hospital de Caridade",               "Jaguaruna",        "R$ 400.000",   "EM ANÁLISE"),
    ("16", "Creche Ouro Negro (CEI)",            "Forquilhinha",     "R$ 900.000",   "AG. PORTARIA"),
    ("17", "Cobertura de Cancha",                "Paulo Lopes",      "R$ 900.000",   "PUBLICADO"),
    ("18", "Ginásio + Pavimentação",             "Paulo Lopes",      "R$ 600.000",   "PRONTO"),
    ("19", "Viatura Corpo de Bombeiros",         "Tubarão",          "R$ 360.000",   "PAGO"),
    ("20", "Câmeras de Segurança",               "Biguaçu",          "R$ 250.000",   "AG. PORTARIA"),
    ("21", "Viatura Delegacia",                  "Laguna",           "R$ 150.000",   "PRONTO"),
    ("22", "Drone Agrícola",                     "Treze de Maio",    "R$ 100.000",   "AG. PORTARIA"),
    ("23", "2 Tratores Agrícolas",               "Braço do Norte",   "R$ 400.000",   "PRONTO"),
]

tbl = doc.add_table(rows=1+len(rows_data), cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
col_w = [Cm(1), Cm(5.5), Cm(3.5), Cm(2.8), Cm(3)]

# Header
for j, h in enumerate(headers):
    tbl.cell(0, j).width = col_w[j]
    set_cell_bg(tbl.cell(0, j), AZUL_ESCURO)
    p = tbl.cell(0, j).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BRANCO; r.font.name = "Calibri"

for i, row_d in enumerate(rows_data):
    bg = CINZA_SUAVE if i % 2 == 0 else BRANCO
    for j, val in enumerate(row_d):
        cell = tbl.cell(i+1, j)
        cell.width = col_w[j]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in [0,3,4] else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent  = Cm(0) if j in [0,3,4] else Cm(0.2)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        color = CINZA_TEXTO
        bold  = False
        if j == 3: color = AZUL_MEDIO; bold = True
        if j == 4:
            if val == "PAGO":      color = VERDE
            elif val == "PRONTO":  color = VERDE
            elif "ANÁLISE" in val: color = LARANJA
            bold = True
        r = p.add_run(val)
        r.font.size = Pt(8.5); r.font.color.rgb = color; r.bold = bold; r.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ─── RODAPÉ FINAL ──────────────────────────────────────────────────────────────
tbl_rod = doc.add_table(rows=1, cols=1)
c_rod = tbl_rod.cell(0,0)
set_cell_bg(c_rod, AZUL_ESCURO)
p_rod = c_rod.paragraphs[0]
p_rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_rod.paragraph_format.space_before = Pt(8)
p_rod.paragraph_format.space_after  = Pt(8)
r_rod = p_rod.add_run("Mandato Dep. Pepê Collaço  |  Progressistas  |  20ª Legislatura — 2023/2026  |  Sul Catarinense")
r_rod.font.size = Pt(9); r_rod.font.color.rgb = RGBColor(0xAA,0xC4,0xE8); r_rod.font.name = "Calibri"

# ─── SALVAR ─────────────────────────────────────────────────────────────────────
output = "/Users/luisgustavodebiasi/TRABALHOS/Projetos Externo/PEPE/REDES/Emendas_Conteúdo_Redes_Sociais.docx"
doc.save(output)
print(f"Salvo em: {output}")
