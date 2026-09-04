# -*- coding: utf-8 -*-
"""
Base SIMPLES para os roteiros de vídeo do Mandato Pepê Collaço.
Sem tabelas, sem layout colorido. Cada bloco = IMAGENS (o que aparece) e,
logo abaixo, FALA (o que se diz). Legenda do post no final.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PRETO = RGBColor(0x1A, 0x1A, 0x1A)
AZUL  = RGBColor(0x1A, 0x37, 0x6C)
CINZA = RGBColor(0x60, 0x60, 0x60)
FONT  = "Calibri"

OUT = os.path.dirname(os.path.abspath(__file__))


def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(11)
    st.font.color.rgb = PRETO
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    return doc


def _p(doc, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def titulo(doc, texto, sub=""):
    p = _p(doc, after=2)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = AZUL
    r.font.name = FONT
    if sub:
        ps = _p(doc, after=10)
        rs = ps.add_run(sub)
        rs.italic = True
        rs.font.size = Pt(10.5)
        rs.font.color.rgb = CINZA
        rs.font.name = FONT


def secao(doc, texto):
    """Cabeçalho de seção interno (ex.: para vídeos com vários sub-roteiros)."""
    p = _p(doc, after=4, before=10)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = AZUL
    r.font.name = FONT


def bloco(doc, imagens, fala):
    """Um beat do vídeo: IMAGENS em cima, FALA embaixo."""
    pi = _p(doc, after=2, before=6)
    ri = pi.add_run("IMAGENS: ")
    ri.bold = True
    ri.font.color.rgb = AZUL
    ri.font.size = Pt(11)
    ri2 = pi.add_run(imagens)
    ri2.font.size = Pt(11)
    ri2.font.color.rgb = PRETO

    pf = _p(doc, after=8)
    rf = pf.add_run("FALA: ")
    rf.bold = True
    rf.font.color.rgb = AZUL
    rf.font.size = Pt(11)
    if fala:
        rf2 = pf.add_run(f"“{fala}”")
        rf2.font.size = Pt(11)
        rf2.font.color.rgb = PRETO
    else:
        rf2 = pf.add_run("(sem fala — só trilha e texto na tela)")
        rf2.italic = True
        rf2.font.size = Pt(10.5)
        rf2.font.color.rgb = CINZA


def texto(doc, t, italic=False, bold=False, after=6):
    p = _p(doc, after=after)
    r = p.add_run(t)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(10.5)
    r.font.color.rgb = CINZA if italic else PRETO
    r.font.name = FONT


def linha(doc):
    p = _p(doc, after=6, before=6)
    r = p.add_run("— — —")
    r.font.color.rgb = CINZA


def legenda(doc, txt, hashtags=""):
    p = _p(doc, after=4, before=12)
    r = p.add_run("LEGENDA DO POST")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = AZUL
    r.font.name = FONT
    pt = _p(doc, after=4)
    rt = pt.add_run(txt)
    rt.font.size = Pt(11)
    rt.font.color.rgb = PRETO
    if hashtags:
        ph = _p(doc, after=6)
        rh = ph.add_run(hashtags)
        rh.font.size = Pt(10.5)
        rh.font.color.rgb = AZUL


def aviso(doc, itens):
    p = _p(doc, after=4, before=12)
    r = p.add_run("⚠️ Conferir antes de publicar")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = CINZA
    for it in itens:
        pi = _p(doc, after=2)
        pi.paragraph_format.left_indent = Cm(0.4)
        ri = pi.add_run("• " + it)
        ri.font.size = Pt(10)
        ri.font.color.rgb = CINZA


def salvar(doc, nome):
    doc.save(os.path.join(OUT, nome))
    print("OK:", nome)
